import streamlit as st
from groq import Groq
import os
import re
import random
from docx import Document
from docx.shared import RGBColor
import sys   # 🔹 thêm dòng này

# =========================
# 📁 Hàm xử lý đường dẫn (hoạt động cho cả .py và .exe)
# =========================
def resource_path(relative_path):
    """Trả về đường dẫn tuyệt đối tới resource (data, file, v.v.)
    Dùng cho cả khi chạy Python và khi chạy bản đóng gói PyInstaller.
    """
    try:
        base_path = sys._MEIPASS  # thư mục tạm khi đóng gói .exe
    except Exception:
        base_path = os.path.abspath(".")  # khi chạy .py
    return os.path.join(base_path, relative_path)

# =========================
# ⚙️ Cấu hình trang
# =========================
st.set_page_config(layout="wide")
st.title("📝 Sinh đề kiểm tra từ ma trận (chuẩn ex_test)")

# =========================
# 🔑 Nhập API Key
# =========================
api_key = st.text_input("Nhập API Key của Groq:", type="password")

# =========================
# 🧠 Hàm tiện ích
# =========================
def get_sample_file(lop, topic, dang_cauhoi, muc_do, dang):
    base_dir = resource_path("data")
    folder = os.path.join(base_dir, lop, topic, dang_cauhoi, muc_do)
    filename = f"{dang}.txt"
    return os.path.join(folder, filename)

def split_ex_blocks(latex_text):
    """Tách từng câu hỏi \\begin{ex} ... \\end{ex}"""
    return re.findall(r"\\begin{ex}.*?\\end{ex}", latex_text, re.S)

# =========================
# 💾 Xuất LaTeX
# =========================
def export_latex_ex(all_questions, filename="output.tex"):
    latex_content = (
        "\\documentclass[12pt]{article}\n"
        "\\usepackage[utf8]{vietnam}\n"
        "\\usepackage{ex_test}\n"
        "\\begin{document}\n"
        "\\section*{Đề kiểm tra}\n"
    )
    latex_content += "\n\n".join(all_questions)
    latex_content += "\n\\end{document}"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(latex_content)
    return filename

# =========================
# 💾 Xuất Word
# =========================
def export_word_ex(all_questions, filename="output.docx"):
    doc = Document()
    doc.add_heading("Đề kiểm tra", 0)
    questions = []
    for q in all_questions:
        questions.extend(split_ex_blocks(q))

    for i, q in enumerate(questions, 1):
        noi_dung_match = re.search(
            r"\\begin{ex}(.*?)(?=\\choice|\\choiceTF|\\shortans|\\loigiai|\\end{ex})", q, re.S
        )
        noi_dung = noi_dung_match.group(1).strip() if noi_dung_match else q
        p = doc.add_paragraph()
        run_q = p.add_run(f"Câu {i}. ")
        run_q.bold = True
        p.add_run(noi_dung)

        dap_an = None

        # --- Trắc nghiệm nhiều lựa chọn ---
        if "\\choice" in q and not "\\choiceTF" in q:
            lc_block = re.search(r"\\choice(.*?)(?=\\loigiai|\\end{ex})", q, re.S)
            if lc_block:
                lines = lc_block.group(1).splitlines()
                options = []
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    is_true = "\\True" in line
                    line = line.replace("\\True", "").strip("{} ")
                    options.append((line, is_true))
                for j, (opt, is_true) in enumerate(options):
                    label = chr(65+j) + "."
                    p = doc.add_paragraph()
                    run = p.add_run(f"{label} {opt}")
                    if is_true:
                        run.bold = True
                        run.underline = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        dap_an = chr(65+j)

        # --- Đúng/Sai ---
        elif "\\choiceTF" in q:
            tf_block = re.search(r"\\choiceTF(.*?)(?=\\loigiai|\\end{ex})", q, re.S)
            if tf_block:
                lines = tf_block.group(1).splitlines()
                tf_ans = ""
                idx_tf = 0
                for line in lines:
                    line = line.strip("{} \t")
                    if not line:
                        continue
                    is_true = "\\True" in line
                    clean_line = line.replace("\\True", "").strip()
                    label = f"{chr(97+idx_tf)})"
                    p = doc.add_paragraph()
                    run = p.add_run(f"{label} {clean_line}")
                    if is_true:
                        run.bold = True
                        run.underline = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    tf_ans += "Đ" if is_true else "S"
                    idx_tf += 1
                dap_an = tf_ans

        # --- Trả lời ngắn ---
        elif "\\shortans" in q:
            sa_block = re.search(r"\\shortans\{(.*?)\}", q)
            if sa_block:
                doc.add_paragraph("Trả lời ngắn: ............")
                dap_an = sa_block.group(1).strip()

        # --- Lời giải ---
        loi_giai_match = re.search(r"\\loigiai\{(.*?)\}", q, re.S)
        if loi_giai_match:
            loi_giai = loi_giai_match.group(1).strip()
            p = doc.add_paragraph()
            run_lg = p.add_run("Lời giải: ")
            run_lg.bold = True
            if dap_an:
                p.add_run(f"Đáp án: {dap_an}. {loi_giai}")
            else:
                p.add_run(loi_giai)
        else:
            if dap_an:
                p = doc.add_paragraph()
                run_lg = p.add_run("Lời giải: ")
                run_lg.bold = True
                p.add_run(f"Đáp án: {dap_an}.")
    doc.save(filename)
    return filename

# =========================
# 📂 Đọc danh sách thư mục động
# =========================
def list_subfolders(path):
    if not os.path.exists(path):
        return []
    return sorted([f for f in os.listdir(path) if os.path.isdir(os.path.join(path, f))])

def list_txt_files(path):
    if not os.path.exists(path):
        return []
    return sorted([f[:-4] for f in os.listdir(path) if f.endswith(".txt")])

# =========================
# 🧩 Giao diện chọn ma trận câu hỏi
# =========================
st.markdown("## 🧩 Ma trận chọn câu hỏi")

BASE_DIR = "data"
ALL_LOPS = list_subfolders(BASE_DIR)

if "configs" not in st.session_state:
    st.session_state.configs = [{"lop": "", "topic": "", "dang_cauhoi": "", "muc_do": "", "dang": "", "count": 1}]

if st.button("➕ Thêm cấu hình"):
    st.session_state.configs.append({"lop": "", "topic": "", "dang_cauhoi": "", "muc_do": "", "dang": "", "count": 1})
    st.rerun()

for idx, cfg in enumerate(list(st.session_state.configs)):
    cols = st.columns([1.2, 1.6, 1.4, 1.4, 1.6, 0.9, 0.8])

    # --- Lớp ---
    with cols[0]:
        if ALL_LOPS:
            current_lop = cfg.get("lop", "")
            if current_lop not in ALL_LOPS:
                current_lop = ALL_LOPS[0]
            cfg["lop"] = st.selectbox("Lớp", ALL_LOPS, index=ALL_LOPS.index(current_lop), key=f"lop_{idx}")
        else:
            st.warning("⚠️ Thư mục data chưa có lớp nào.")
            cfg["lop"] = ""

    # --- Chủ đề ---
    topics = list_subfolders(os.path.join(BASE_DIR, cfg["lop"])) if cfg["lop"] else []
    with cols[1]:
        if topics:
            current_topic = cfg.get("topic", "")
            if current_topic not in topics:
                current_topic = topics[0]
            cfg["topic"] = st.selectbox("Chủ đề", topics, index=topics.index(current_topic), key=f"topic_{idx}")
        else:
            st.text_input("Chủ đề", value="(trống)", key=f"topic_{idx}_empty", disabled=True)
            cfg["topic"] = ""

    # --- Loại câu hỏi ---
    dang_cauhoi_list = list_subfolders(os.path.join(BASE_DIR, cfg["lop"], cfg["topic"])) if cfg["topic"] else []
    with cols[2]:
        if dang_cauhoi_list:
            current_dang_cauhoi = cfg.get("dang_cauhoi", "")
            if current_dang_cauhoi not in dang_cauhoi_list:
                current_dang_cauhoi = dang_cauhoi_list[0]
            cfg["dang_cauhoi"] = st.selectbox("Loại", dang_cauhoi_list, index=dang_cauhoi_list.index(current_dang_cauhoi), key=f"dangcauhoi_{idx}")
        else:
            st.text_input("Loại", value="(trống)", key=f"dangcauhoi_{idx}_empty", disabled=True)
            cfg["dang_cauhoi"] = ""

    # --- Mức độ ---
    mucdo_list = list_subfolders(os.path.join(BASE_DIR, cfg["lop"], cfg["topic"], cfg["dang_cauhoi"])) if cfg["dang_cauhoi"] else []
    with cols[3]:
        if mucdo_list:
            current_mucdo = cfg.get("muc_do", "")
            if current_mucdo not in mucdo_list:
                current_mucdo = mucdo_list[0]
            cfg["muc_do"] = st.selectbox("Mức độ", mucdo_list, index=mucdo_list.index(current_mucdo), key=f"mucdo_{idx}")
        else:
            st.text_input("Mức độ", value="(trống)", key=f"mucdo_{idx}_empty", disabled=True)
            cfg["muc_do"] = ""

    # --- Dạng (file .txt) ---
    dang_files = list_txt_files(os.path.join(BASE_DIR, cfg["lop"], cfg["topic"], cfg["dang_cauhoi"], cfg["muc_do"])) if cfg["muc_do"] else []
    with cols[4]:
        if dang_files:
            current_dang = cfg.get("dang", "")
            if current_dang not in dang_files:
                current_dang = dang_files[0]
            cfg["dang"] = st.selectbox("Dạng", dang_files, index=dang_files.index(current_dang), key=f"dang_{idx}")
        else:
            st.text_input("Dạng", value="(trống)", key=f"dang_{idx}_empty", disabled=True)
            cfg["dang"] = ""

    # --- Số lượng ---
    with cols[5]:
        cfg["count"] = st.number_input("Số lượng", min_value=1, max_value=50, value=cfg.get("count", 1), key=f"count_{idx}")

    # --- Xóa ---
    with cols[6]:
        if st.button("❌", key=f"remove_{idx}"):
            st.session_state.configs.pop(idx)
            st.rerun()

tong_cau = sum(c.get("count", 0) for c in st.session_state.configs)
st.info(f"📊 Tổng số câu hỏi đã chọn: {tong_cau}")

# =========================
# 🚀 Sinh câu hỏi từ ma trận
# =========================
col_gen = st.columns([1,1,1])
with col_gen[0]:
    submitted = st.button("🚀 Sinh câu hỏi từ ma trận")
with col_gen[1]:
    export_word_btn = st.button("⬇️ Xuất Word")
with col_gen[2]:
    export_tex_btn = st.button("⬇️ Xuất LaTeX")

if "all_questions" not in st.session_state:
    st.session_state.all_questions = []

if submitted:
    if not api_key:
        st.error("Vui lòng nhập API Key trước khi sinh câu hỏi.")
    else:
        client = Groq(api_key=api_key)
        all_questions = []
        for idx, cfg in enumerate(st.session_state.configs):
            file_path = get_sample_file(cfg["lop"], cfg["topic"], cfg["dang_cauhoi"], cfg["muc_do"], cfg["dang"])
            if not os.path.exists(file_path):
                st.warning(f"❌ Không tìm thấy file: {file_path}")
                continue

            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Nếu là file chứa sẵn ex
            if cfg["dang"].endswith("_file"):
                ex_blocks = split_ex_blocks(content)
                if not ex_blocks:
                    st.warning(f"❌ File {cfg['dang']}.txt không có câu hỏi nào.")
                    continue
                selected = random.sample(ex_blocks, min(cfg["count"], len(ex_blocks)))
                all_questions.extend(selected)
                st.success(f"✅ Đã lấy {len(selected)} câu từ file.")
                continue

            if "\\choiceTF" in content:
                cau_truc = "luôn dùng \\choiceTF"
            elif "\\shortans" in content:
                cau_truc = "luôn dùng \\shortans"
            else:
                cau_truc = "luôn dùng \\choice"

            prompt = f"""
Đây là các câu hỏi mẫu theo chuẩn gói ex_test:

{content}

Hãy sinh {cfg['count']} câu hỏi tương tự bằng tiếng Việt.
Yêu cầu:
- Dùng môi trường \\begin{{ex}} ... \\end{{ex}}
- {cau_truc}
- Mỗi câu có \\loigiai{{...}} ở cuối
- Nếu có hình tikz thì sinh code tikz phù hợp

⚠️ Chỉ trả về LaTeX, không thêm chữ nào khác.
"""
            try:
                chat_completion = client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model="llama-3.3-70b-versatile",
                    temperature=0.7,
                )
                output = chat_completion.choices[0].message.content.strip()
                st.code(output, language="latex")
                all_questions.append(output)
                st.success(f"✅ Sinh thành công {cfg['count']} câu.")
            except Exception as e:
                st.error(f"Lỗi khi gọi Groq API: {e}")

        st.session_state.all_questions = all_questions
        st.success(f"🎯 Hoàn tất sinh đề: {len(all_questions)} câu.")

# =========================
# 💾 Xuất file
# =========================
if export_word_btn:
    if not st.session_state.all_questions:
        st.warning("Chưa có câu hỏi để xuất.")
    else:
        word_file = export_word_ex(st.session_state.all_questions, "de_kiem_tra.docx")
        with open(word_file, "rb") as f:
            st.download_button("⬇️ Tải Word", f, file_name="de_kiem_tra.docx")

if export_tex_btn:
    if not st.session_state.all_questions:
        st.warning("Chưa có câu hỏi để xuất.")
    else:
        tex_file = export_latex_ex(st.session_state.all_questions, "de_kiem_tra.tex")
        with open(tex_file, "rb") as f:
            st.download_button("⬇️ Tải LaTeX", f, file_name="de_kiem_tra.tex")

# =========================
# 👀 Preview
# =========================
if st.session_state.all_questions:
    st.markdown("### Xem trước (5 câu đầu)")
    for q in st.session_state.all_questions[:5]:
        st.code(q, language="latex")
